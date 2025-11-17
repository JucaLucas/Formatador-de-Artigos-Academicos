from flask import Flask, request, send_file
from flask_cors import CORS
from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_COLOR_INDEX
from docx.shared import Pt, Cm, RGBColor
from docx.oxml.ns import qn
import tempfile
import re
from Condicoes import cidades_bahia, palavras_orientador, cursos, instituicao
import unicodedata


app = Flask(__name__)
CORS(app)


ja_tem_autor = False
ja_tem_cidade = False
ja_tem_instituicao = False
ja_tem_titulo_capa = False
ja_tem_curso = False
ja_tem_ano = False


def _normalizar_para_busca(texto):
    """Remove pontuação chata, sufixos como '- BA' ou '/BA', múltiplos espaços, e deixa minusculo."""
    if not texto:
        return ""
    t = texto.strip().lower()

    # remove sufixos comuns: "- ba", "/ba", ", ba", " - bahia", "/bahia", etc.
    t = re.sub(r'[-/]\s*ba(hia)?\b', '', t)            # "- ba", "/bahia"
    t = re.sub(r',\s*ba(hia)?\b', '', t)               # ", BA"
    t = re.sub(r'\s+ba(hia)?\b', '', t)                # " BA" ou " BAHIA"
    t = re.sub(r'\bestado da bahia\b', '', t)

    # remove pontuação (mantém letras e espaços)
    t = re.sub(r'[^0-9a-zà-ú\s]', ' ', t, flags=re.IGNORECASE)

    # colapsa espaços e trim
    t = re.sub(r'\s+', ' ', t).strip()
    return t

def detectar_tipo(linha, identificados):
    texto_original = linha.strip()
    texto_lower = texto_original.lower()

    # ============================================
    # 1) ANO (regra mais simples e 100% segura)
    # ============================================
    if texto_original.isdigit() and len(texto_original) == 4:
        return "ano"

    # ============================================
    # 2) CIDADE (com lista prévia e UF opcional)
    # ============================================
    if identificados["cidade"] is None:
        cidade = eh_cidade(texto_original, cidades_bahia)
        if cidade:
            return "cidade"

    # 3) INSTITUIÇÃO (função dedicada)
    if identificados["instituicao"] is None:
        if eh_instituicao(texto_original):
            return "instituicao"


    # ============================================
    # 4) CURSO / TIPO DE TRABALHO
    # ============================================

    if identificados["curso"] is None:

        texto_norm = remover_acentos(texto_lower)
        for c in cursos:
            if texto_norm == remover_acentos(c.lower()):
                return "curso"

    if "curso de " in texto_norm or "bacharelado em" in texto_norm:
        if len(texto_original.split()) >= 2:
            return "curso"

    # ============================================
    # 5) AUTOR (regras MUITO fortes)
    # 2 a 4 nomes, iniciando com maiúscula, sem números
    # ============================================
    if identificados["autor"] is None:
        palavras = texto_original.split()

        # 2 a 4 palavras
        if 2 <= len(palavras) <= 4:
            # Todas começam com Maiúscula
            if all(p[0].isupper() and p[1:].islower() for p in palavras if len(p) > 1):
                # Não contém palavras de instituição
                if not any(k in texto_lower for k in instituicao):
                    # Não contém curso
                    if not any(k in texto_lower for k in cursos):
                        # Não contém cidade
                        if not eh_cidade(texto_original, cidades_bahia):
                            return "autor"

    # ============================================
    # 6) TÍTULO (≥ 40 caracteres e ≥ 6 palavras)
    # ============================================
    if identificados["titulo"] is None:
        if len(texto_original) >= 40 and len(texto_original.split()) >= 6:
            return "titulo"

    # ============================================
    # 7) SUBTÍTULO (só pode vir após título)
    # ============================================
    if identificados["titulo"] is not None and identificados["subtitulo"] is None:
        if 10 <= len(texto_original) <= 80:
            return "subtitulo"

    # ============================================
    # Nada identificado
    # ============================================
    return None


def identificar_linhas_da_capa(doc):
    linhas = [p.text for p in doc.paragraphs]

    identificados = {
        "instituicao": None,
        "autor": None,
        "curso": None,
        "titulo": None,
        "subtitulo": None,
        "cidade": None,
        "ano": None
    }

    classificados = set()  # <- impede alterar classificação depois

    for _ in range(8):  # repete para refinar a identificação
        for i, linha in enumerate(linhas):
            if not linha.strip():
                continue

            if i in classificados:
                continue  # <- já foi classificado, não muda mais

            tipo = detectar_tipo(linha, identificados)

            if tipo and identificados[tipo] is None:
                identificados[tipo] = i
                classificados.add(i)  # <- trava a linha como aquele tipo

    return identificados


def detectar_fonte_principal(doc):
    contagem_fontes = {}

    # Verifica a fonte configurada no estilo Normal
    try:
        estilo_normal = doc.styles["Normal"].font.name
    except:
        estilo_normal = None

    for p in doc.paragraphs:
        for run in p.runs:
            fonte = run.font.name or estilo_normal or "Arial"
            fonte = fonte.strip()

            if fonte not in contagem_fontes:
                contagem_fontes[fonte] = 0

            contagem_fontes[fonte] += len(run.text)

    # Descobre a fonte mais usada
    fonte_predominante = max(contagem_fontes, key=contagem_fontes.get)

    # Normaliza para Arial ou Times
    if "arial" in fonte_predominante.lower():
        return "Arial"

    if "times" in fonte_predominante.lower():
        return "Times New Roman"

    # Se for outra fonte → força Arial
    return "Arial"

def classificar_capa(texto):
    texto_limpo = texto.strip()

    if not texto_limpo:
        return "vazio"

    # Instituição: normalmente grande, não começa com número e tem várias palavras
    if len(texto_limpo.split()) >= 3 and texto_limpo.isupper():
        return "instituicao"

    # Autor: costuma ter nome + sobrenome iniciando maiúsculas
    if re.match(r'^[A-ZÁÀÃÂÉÍÓÚ][a-z].+ [A-ZÁÀÃÂÉÍÓÚ][a-z].+', texto_limpo):
        return "autor"

    # Curso
    if any(palavra in texto_limpo.lower() for palavra in ["curso", "bacharel", "licenciatura", "tecnólogo"]):
        return "curso"

    # Título da capa: mais longo, não precisa ser uppercase
    if len(texto_limpo.split()) >= 4 and not re.match(r'^\d', texto_limpo):
        return "titulo_capa"

    # Subtitulo (se houver)
    if texto_limpo.lower().startswith(("um estudo", "análise", "desenvolvimento", "projeto", "monografia", "trabalho")):
        return "subtitulo"

    # Cidade
    if texto_limpo.lower() in ["feira de santana", "salvador", "são paulo", "rio de janeiro"]:
        return "cidade"

    # Ano (4 dígitos)
    if re.match(r'^\d{4}$', texto_limpo):
        return "ano"

    return "outro"

def eh_ano(texto):
    texto = (texto or "").strip()
    # procura um ano isolado como 1999, 2025, 2010 (entre limites razoáveis)
    m = re.search(r"\b(19|20)\d{2}\b", texto)
    if m:
        # opcional: retorna o ano encontrado se precisar
        return True
    return False



from Condicoes import cidades_bahia


def remover_acentos(txt):
    """Remove acentos e normaliza texto para comparação segura."""
    return ''.join(
        c for c in unicodedata.normalize('NFD', txt)
        if unicodedata.category(c) != 'Mn'
    )

def eh_cidade(texto):
    texto_norm = remover_acentos(texto.strip().lower())

    # CIDADES EXATAS
    for c in cidades_bahia:
        if remover_acentos(c.lower()) == texto_norm:
            return True

    # CIDADE + UF (ex: Salvador - BA, Feira de Santana BA)
    if texto_norm.endswith(" ba") or texto_norm.endswith("-ba"):
        return True

    return False

print(eh_cidade("Feira de Santana"))

def eh_instituicao(texto):
    texto_lower = texto.lower()

    for inst in instituicao:
        inst_lower = inst.lower().strip()

        # Evita falso positivo parcial (ex: "centro" dentro de "concentração")
        if f" {inst_lower} " in f" {texto_lower} ":
            return True

        if texto_lower.startswith(inst_lower):
            return True

    return False


def eh_curso(txt): 
    t = remover_acentos(txt.strip().lower()) 
    for c in cursos: 
        if remover_acentos(c.lower()) in t: 
            return True
    return False

def eh_autor(t):
    # Nome de pessoa: 2 a 5 palavras com iniciais maiúsculas
    return re.fullmatch(r"([A-ZÁÉÍÓÚÂÊÔÃÕ][a-záéíóúâêôãõç]+)(\s+[A-ZÁÉÍÓÚÂÊÔÃÕ][a-záéíóúâêôãõç]+){1,4}", t) is not None

def classificar_linhas(linhas):
    """
    Recebe lista de linhas (texto bruto) e retorna dict com os primeiros
    valores detectados para: ano, instituicao, titulo, subtitulo, autor, cidade, curso
    """
    identificados = {
        "ano": None,
        "instituicao": None,
        "titulo": None,
        "subtitulo": None,
        "autor": None,
        "cidade": None,
        "curso": None
    }

    titulo_identificado = False
    resto = linhas[:]

    for _ in range(7):  # repete para refinar
        novos_resto = []
        for t in resto:
            s = t.strip()
            if not s:
                continue

            # 1 — ANO
            if identificados["ano"] is None and eh_ano(s):
                identificados["ano"] = s
                continue

            # 2 — INSTITUIÇÃO
            if identificados["instituicao"] is None and eh_instituicao(s):
                identificados["instituicao"] = s
                continue

            # 3 — CURSO (antes do autor para evitar confusão)
            if identificados["curso"] is None and eh_curso(s):
                identificados["curso"] = s
                continue

            # 4 — TÍTULO (usa estado)
            if identificados["titulo"] is None:
                eh, novo_estado = eh_titulo(s, titulo_identificado)
                if eh:
                    identificados["titulo"] = s
                    titulo_identificado = novo_estado
                    continue

            # 5 — SUBTÍTULO (só após título)
            if identificados["titulo"] and identificados["subtitulo"] is None:
                if eh_subtitulo(s, titulo_identificado):
                    identificados["subtitulo"] = s
                    continue

            # 6 — AUTOR
            if identificados["autor"] is None and eh_autor(s):
                identificados["autor"] = s
                continue

            # 7 — CIDADE
            if identificados["cidade"] is None and eh_cidade(s):
                identificados["cidade"] = s
                continue

            novos_resto.append(t)

        if len(novos_resto) == len(resto):
            break
        resto = novos_resto

    return identificados

def normalizar(texto):
    if not texto:
        return ""
    texto = texto.strip().lower()
    texto = unicodedata.normalize('NFD', texto)
    texto = "".join(c for c in texto if unicodedata.category(c) != "Mn")
    texto = " ".join(texto.split())  # remove espaços duplicados
    return texto


def formatar_capa(doc):
    """
    Analisa os parágrafos após o resumo (sem formatar nada) e depois
    reconstrói apenas a capa na ordem ABNT.

    A análise verifica:
    - Parágrafos devem ser JUSTIFICADOS
    - Títulos de seções devem ser alinhados À ESQUERDA
    """

    avisos = []  # lista que será retornada ao final

    # ---------------------------------------------------------
    # 1) ENCONTRAR O RESUMO
    # ---------------------------------------------------------
    index_resumo = None
    for idx, p in enumerate(doc.paragraphs):
        if p.text and p.text.strip().lower().startswith("resumo"):
            index_resumo = idx
            break

    limite = index_resumo if index_resumo is not None else len(doc.paragraphs)

    # ---------------------------------------------------------
    # 2) FAZER ANÁLISE ABNT *ANTES* DE FORMATAR A CAPA
    # ---------------------------------------------------------
    if index_resumo is not None:
        for idx in range(index_resumo + 1, len(doc.paragraphs)):
            p = doc.paragraphs[idx]
            texto = p.text.strip()

            if not texto:
                continue

            # Detecta SEÇÕES — ABNT exige alinhamento à esquerda
            if texto.isupper() and len(texto.split()) <= 6:
                if p.alignment != WD_ALIGN_PARAGRAPH.LEFT:
                    avisos.append(f"A seção '{texto}' deve ser alinhada à esquerda.")
                continue

            # ANALISAR PARÁGRAFOS COMUNS (devem ser justificados)
            if p.alignment != WD_ALIGN_PARAGRAPH.JUSTIFY:
                avisos.append(f"O parágrafo '{texto[:40]}...' deve ser justificado.")

    # ---------------------------------------------------------
    # 3) CAPTURA DA CAPA (PARTE EXISTENTE DO SEU CÓDIGO)
    # ---------------------------------------------------------
    linhas = [p.text for p in doc.paragraphs[:limite]]

    try:
        identificados_indices = identificar_linhas_da_capa(doc)
    except Exception:
        identificados_indices = {}
        classificados = classificar_linhas(linhas)
        for chave, valor in classificados.items():
            if valor:
                try:
                    identificados_indices[chave] = linhas.index(valor)
                except ValueError:
                    identificados_indices[chave] = None
            else:
                identificados_indices[chave] = None

    def txt(idx):
        return linhas[idx].strip() if (idx is not None and 0 <= idx < len(linhas) and linhas[idx].strip()) else None

    inst_txt = txt(identificados_indices.get("instituicao"))
    curso_txt = txt(identificados_indices.get("curso"))
    autor_txt = txt(identificados_indices.get("autor"))
    titulo_txt = txt(identificados_indices.get("titulo"))
    subt_txt = txt(identificados_indices.get("subtitulo"))
    cidade_txt = txt(identificados_indices.get("cidade"))
    ano_txt = txt(identificados_indices.get("ano"))

    # fallbacks
    if not titulo_txt:
        for l in linhas:
            if l and eh_titulo(l, False)[0]:
                titulo_txt = l.strip()
                break

    if not autor_txt:
        for l in linhas:
            if l and eh_autor(l):
                autor_txt = l.strip()
                break

    if not ano_txt:
        for l in linhas:
            if l and eh_ano(l):
                ano_txt = l.strip()
                break

    if not cidade_txt:
        for l in linhas:
            if l and eh_cidade(l):
                cidade_txt = l.strip()
                break

    # monta ordem
    capa_ordem = []
    if inst_txt:
        capa_ordem.append(("instituicao", inst_txt))
    if curso_txt:
        capa_ordem.append(("curso", curso_txt))
    if autor_txt:
        capa_ordem.append(("autor", autor_txt))
    if titulo_txt:
        capa_ordem.append(("titulo_capa", titulo_txt))
    if subt_txt and subt_txt != titulo_txt:
        capa_ordem.append(("subtitulo", subt_txt))
    if cidade_txt:
        capa_ordem.append(("cidade", cidade_txt))
    if ano_txt:
        capa_ordem.append(("ano", ano_txt))

    novos_pares = []
    usados = set()
    for tipo, texto in capa_ordem:
        if texto not in usados:
            novos_pares.append((tipo, texto))
            usados.add(texto)

    capa_ordem = novos_pares

    # ---------------------------------------------------------
    # 4) FORMATAÇÃO DA CAPA (NÃO ALTERADA)
    # ---------------------------------------------------------
    first_idx = 0
    for idx in range(limite):
        if doc.paragraphs[idx].text.strip():
            first_idx = idx
            break

    for offset, (tipo, texto) in enumerate(capa_ordem):
        target_idx = first_idx + offset
        p = doc.paragraphs[target_idx]
        p.text = texto.strip()
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        p.paragraph_format.right_indent = Pt(0)


        for run in p.runs:
            run.font.size = Pt(12)
            run.bold = False

        if tipo == "titulo_capa":
            original = p.text
            p.text = ""
            run = p.add_run(original.upper())
            run.font.size = Pt(14)
            run.bold = True

    start_clean = first_idx + len(capa_ordem)
    for idx in range(start_clean, limite):
        doc.paragraphs[idx].text = ""

    # ---------------------------------------------------------
    # 5) RETORNA OS AVISOS AO USUÁRIO
    # ---------------------------------------------------------
    return avisos


# -----------------------------------------------------------
# 🔹 Estado global para garantir que o título só apareça 1x
# -----------------------------------------------------------



from docx import Document

def eh_titulo(texto, titulo_identificado):
    """
    Detecta se o texto é o título principal da capa.
    Retorna uma tupla: (eh_titulo, novo_estado)
    """
    texto_limpo = texto.strip()
    if not texto_limpo or titulo_identificado:
        return False, titulo_identificado  # já identificou antes → não é título

    # Critérios mais fortes:
    # 1. Todo em maiúsculo (ou quase)
    # 2. Comprimento mínimo de 4 palavras
    maiusculas = sum(1 for c in texto_limpo if c.isupper())
    minusculas = sum(1 for c in texto_limpo if c.islower())
    proporcao_maiusculas = maiusculas / max(1, (maiusculas + minusculas))

    if (proporcao_maiusculas >= 0.85 or texto_limpo.isupper()) and len(texto_limpo.split()) >= 4:
        return True, True  # marcou o título e bloqueia o próximo
    return False, titulo_identificado


def eh_subtitulo(texto, titulo_identificado):
    texto_limpo = texto.strip()
    if not texto_limpo:
        return False

    # Só existe subtítulo se o título já tiver sido identificado
    if not titulo_identificado:
        return False

    # Subtítulo NÃO pode estar totalmente em caixa alta
    if texto_limpo.isupper():
        return False

    # Subtítulo deve ter estrutura de frase (>= 3 palavras)
    if len(texto_limpo.split()) < 3:
        return False

    # Subtítulo deve conter letras minúsculas
    if texto_limpo.upper() == texto_limpo:
        return False

    # Não pode ser nome de autor
    if re.match(r'^[A-ZÁÉÍÓÚÂÊÔÃÕ][a-záéíóúâêôãõç]+(\s+[A-ZÁÉÍÓÚÂÊÔÃÕ][a-záéíóúâêôãõç]+)+$', texto_limpo):
        return False

    # Não pode ser cidade simples (ex: "Feira de Santana")
    if texto_limpo.istitle() and len(texto_limpo.split()) <= 3:
        return False

    # Não pode ser ano
    if re.fullmatch(r'\d{4}', texto_limpo):
        return False

    return True



# 🔹 Exemplo de uso (funciona em qualquer parte do app)
def identificar_titulos(doc):
    titulo_identificado = False

    for p in doc.paragraphs:
        texto = p.text.strip()
        if not texto:
            continue

        eh_tit, titulo_identificado = eh_titulo(texto, titulo_identificado)
        if eh_tit:
            print("Título principal encontrado:", texto)
            continue

        if eh_subtitulo(texto, titulo_identificado):
            print("Subtítulo encontrado:", texto)



# -------- CLASSIFICAR PARÁGRAFOS --------
def classificar_texto(texto, titulo_identificado):
    global ja_tem_autor, ja_tem_cidade, ja_tem_instituicao
    global ja_tem_titulo_capa, ja_tem_curso, ja_tem_ano

    ja_tem_autor = False
    ja_tem_cidade = False
    ja_tem_instituicao = False
    ja_tem_titulo_capa = False
    ja_tem_curso = False
    ja_tem_ano = False
    texto = texto.strip()
    if not texto:
        return "vazio", titulo_identificado

    texto_lower = texto.lower()
    texto_norm = remover_acentos(texto_lower)

    # ============================================================
    # 1 — ANO  (vem antes de tudo)
    # ============================================================
    if not ja_tem_ano and eh_ano(texto):
        ja_tem_ano = True
        return "ano", titulo_identificado

    # ============================================================
    # 2 — CIDADE  (prioridade máxima)
    # ============================================================
    if not ja_tem_cidade and eh_cidade(texto):
        linha_curta = len(texto) <= 40
        termina_com_uf = texto_norm.endswith(" ba") or texto_norm.endswith(" bahia")
        somente_cidade = any(remover_acentos(c.lower()) == texto_norm for c in cidades_bahia)

        if linha_curta or termina_com_uf or somente_cidade:
            ja_tem_cidade = True
            return "cidade", titulo_identificado

    # ============================================================
    # 3 — AUTOR
    # ============================================================
    if not ja_tem_autor and eh_autor(texto):
        ja_tem_autor = True
        return "autor", titulo_identificado

    # ============================================================
    # 4 — ORIENTADOR
    # ============================================================
    if any(p in texto_lower for p in palavras_orientador):
        return "orientador", titulo_identificado

    # ============================================================
    # 5 — INSTITUIÇÃO
    # ============================================================
    if not ja_tem_instituicao and eh_instituicao(texto):
        ja_tem_instituicao = True
        return "instituicao", titulo_identificado

    # ============================================================
    # 6 — CURSO
    # ============================================================
    if not ja_tem_curso and eh_curso(texto):
        ja_tem_curso = True
        return "curso", titulo_identificado

    # ============================================================
    # 7 — TÍTULO DA CAPA
    # ============================================================
    eh_tit, novo_estado = eh_titulo(texto, titulo_identificado)
    if not ja_tem_titulo_capa and eh_tit:
        ja_tem_titulo_capa = True
        return "titulo_capa", novo_estado

    # 8 — SUBTÍTULO DA CAPA (linha logo após o título da capa)
    texto_sem_acentos = remover_acentos(texto)

    if ja_tem_titulo_capa and not ja_tem_autor:
    # Aceita frase maiúscula SEM acentos
        if texto_sem_acentos.isupper() and len(texto.split()) > 1:
          return "subtitulo_capa", titulo_identificado


    # ============================================================
    # 9 — TÍTULOS DO CORPO
    # ============================================================
    if re.match(r"^\d+\s+[A-ZÁÉÍÓÚÂÊÔÃÕ]", texto):
        return "titulo_principal", titulo_identificado

    if re.match(r"^\d+\.\d+(\.\d+)*\s+[A-ZÁÉÍÓÚÂÊÔÃÕ]", texto):
        return "subtitulo", titulo_identificado

    if eh_subtitulo(texto, titulo_identificado):
        return "subtitulo", titulo_identificado

    # ============================================================
    # 10 — PARÁGRAFO NORMAL
    # ============================================================
    return "paragrafo", titulo_identificado


def formatar_resumo(doc):
    """
    Formata o bloco de RESUMO conforme ABNT:
    - Título: RESUMO, centralizado, negrito, fonte 12
    - Texto: justificado, sem recuo, espaçamento 1,5, fonte 12
    """

    mensagens = []

    for i, p in enumerate(doc.paragraphs):
        original = (p.text or "").strip()
        if not original:
            continue

        # Detecta o parágrafo com "RESUMO"
        if original.lower().startswith("resumo"):

            # Conteúdo na mesma linha
            conteudo_mesma_linha = ""
            if ":" in original:
                partes = original.split(":", 1)
                conteudo_mesma_linha = partes[1].strip()

            # --- FORMATAÇÃO DO TÍTULO ---
            p.text = ""
            run = p.add_run("RESUMO")
            run.bold = True
            run.font.size = Pt(12)
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            p.paragraph_format.right_indent = Pt(0)
            p.paragraph_format.left_indent = Pt(0)


            # ===========================
            # 1) Se havia texto após "RESUMO:"
            # ===========================
            if conteudo_mesma_linha:
                if i + 1 < len(doc.paragraphs):
                    novo_para = doc.paragraphs[i + 1].insert_paragraph_before(conteudo_mesma_linha)
                else:
                    novo_para = doc.add_paragraph(conteudo_mesma_linha)

                novo_para.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
                novo_para.paragraph_format.first_line_indent = None
                novo_para.paragraph_format.left_indent = Pt(0)
                novo_para.paragraph_format.right_indent = Pt(0)
                novo_para.paragraph_format.line_spacing = 1.5

                for r in novo_para.runs:
                    r.font.size = Pt(12)

                mensagens.append("✅ Resumo formatado corretamente (conteúdo na mesma linha).")
                return mensagens

            # ===========================
            # 2) Resumo na linha seguinte
            # ===========================
            if i + 1 < len(doc.paragraphs) and doc.paragraphs[i+1].text.strip():
                texto_para = doc.paragraphs[i+1]

                texto_para.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
                texto_para.paragraph_format.first_line_indent = None
                texto_para.paragraph_format.left_indent = None
                texto_para.paragraph_format.right_indent = Pt(0)
                texto_para.paragraph_format.line_spacing = 1.5

                for r in texto_para.runs:
                    r.font.size = Pt(12)

                mensagens.append("✅ Resumo formatado corretamente (linha seguinte).")
                return mensagens

            # ===========================
            # 3) Caso não haja conteúdo
            # ===========================
            if i + 1 < len(doc.paragraphs):
                if not doc.paragraphs[i + 1].text.strip():
                    novo_para = doc.paragraphs[i + 1]
                else:
                    novo_para = doc.paragraphs[i + 1].insert_paragraph_before("")
            else:
                novo_para = doc.add_paragraph("")

            novo_para.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
            novo_para.paragraph_format.first_line_indent = None
            novo_para.paragraph_format.left_indent = None
            novo_para.paragraph_format.right_indent = Pt(0)
            novo_para.paragraph_format.line_spacing = 1.5

            mensagens.append("⚠️ Título RESUMO encontrado, mas o texto estava vazio — criado parágrafo para preenchimento.")
            return mensagens

    mensagens.append("⚠️ Nenhum bloco de RESUMO encontrado.")
    return mensagens

def formatar_abstract(doc):

    for i, p in enumerate(doc.paragraphs):
        original = (p.text or "").strip()
        if not original:
            continue

        # Detecta ABSTRACT
        if original.lower().startswith("abstract"):

            conteudo_mesma_linha = ""
            if ":" in original:
                partes = original.split(":", 1)
                conteudo_mesma_linha = partes[1].strip()

            # --- TÍTULO "ABSTRACT" ---
            p.text = ""
            run = p.add_run("ABSTRACT")
            run.bold = True
            run.font.size = Pt(12)

            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            p.paragraph_format.left_indent = Cm(0)
            p.paragraph_format.first_line_indent = Cm(0)
            p.paragraph_format.right_indent = Pt(0)

            # ============================
            # 1) Texto na mesma linha
            # ============================
            if conteudo_mesma_linha:
                if i + 1 < len(doc.paragraphs):
                    doc.paragraphs[i + 1].insert_paragraph_before(conteudo_mesma_linha)
                    texto_para = doc.paragraphs[i + 1]
                else:
                    texto_para = doc.add_paragraph(conteudo_mesma_linha)

                texto_para.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
                texto_para.paragraph_format.left_indent = None
                texto_para.paragraph_format.first_line_indent = None
                texto_para.paragraph_format.right_indent = Pt(0)
                texto_para.paragraph_format.line_spacing = 1.5

                for r in texto_para.runs:
                    r.font.size = Pt(12)

                return

            # ============================
            # 2) Conteúdo na linha seguinte
            # ============================
            if i + 1 < len(doc.paragraphs) and doc.paragraphs[i + 1].text.strip():
                texto_para = doc.paragraphs[i + 1]

                texto_para.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
                texto_para.paragraph_format.left_indent = None
                texto_para.paragraph_format.first_line_indent = None
                texto_para.paragraph_format.right_indent = Pt(0)
                texto_para.paragraph_format.line_spacing = 1.5

                for r in texto_para.runs:
                    r.font.size = Pt(12)

                return

            # ============================
            # 3) Sem conteúdo → cria vazio
            # ============================
            if i + 1 < len(doc.paragraphs):
                if not doc.paragraphs[i + 1].text.strip():
                    texto_para = doc.paragraphs[i + 1]
                else:
                    texto_para = doc.paragraphs[i + 1].insert_paragraph_before("")
            else:
                texto_para = doc.add_paragraph("")

            texto_para.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
            texto_para.paragraph_format.left_indent = None
            texto_para.paragraph_format.first_line_indent = None
            texto_para.paragraph_format.right_indent = Pt(0)
            texto_para.paragraph_format.line_spacing = 1.5

            return

def formatar_palavras_chave(doc):
    padrao = re.compile(
        r"^(palavras[\s\-]*chaves?|palavraschave)(:)?\s*(.*)$",
        re.IGNORECASE
    )

    for i, p in enumerate(doc.paragraphs):
        texto = p.text.strip()
        if not texto:
            continue

        m = padrao.match(texto)
        if m:
            titulo = "Palavras-chave:"
            conteudo = m.group(3).strip()

            # Remove qualquer recuo
            p.paragraph_format.left_indent = None
            p.paragraph_format.first_line_indent = None
            p.paragraph_format.right_indent = Pt(0)


            if conteudo:
                palavras = re.split(r"[;,.\n]\s*", conteudo)
                palavras = [w.strip() for w in palavras if w.strip()]
                conteudo_formatado = "; ".join(palavras) + "."

                p.text = ""
                run_titulo = p.add_run(titulo)
                run_titulo.bold = True
                p.add_run(" ")
                p.add_run(conteudo_formatado)

                p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
            else:
                p.text = ""
                run_titulo = p.add_run(titulo)
                run_titulo.bold = True
                p.alignment = WD_ALIGN_PARAGRAPH.LEFT

                if i + 1 < len(doc.paragraphs):
                    proximo = doc.paragraphs[i + 1]

                    # Remove recuo do próximo parágrafo também
                    proximo.paragraph_format.left_indent = None
                    proximo.paragraph_format.first_line_indent = None

                    palavras = re.split(r"[;,.\n]\s*", proximo.text.strip())
                    palavras = [w.strip() for w in palavras if w.strip()]
                    proximo.text = "; ".join(palavras) + "."
                    proximo.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY

            # ➜ espaço 1,5 linha antes do ABSTRACT
            p.paragraph_format.space_after = Pt(18)  # ~1,5 linha

            return

def formatar_keywords(doc):
    padrao = re.compile(
        r"^(keywords?)(:)?\s*(.*)$",
        re.IGNORECASE
    )

    for i, p in enumerate(doc.paragraphs):
        texto = p.text.strip()
        if not texto:
            continue

        m = padrao.match(texto)
        if m:
            titulo = "Keywords:"
            conteudo = m.group(3).strip()

            # Remove qualquer recuo
            p.paragraph_format.left_indent = None
            p.paragraph_format.first_line_indent = None
            p.paragraph_format.right_indent = Pt(0)


            if conteudo:
                palavras = re.split(r"[;,.\n]\s*", conteudo)
                palavras = [w.strip() for w in palavras if w.strip()]
                conteudo_formatado = "; ".join(palavras) + "."

                p.text = ""
                p.add_run(titulo).bold = True
                p.add_run(" ")
                p.add_run(conteudo_formatado)
                p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
            else:
                p.text = ""
                p.add_run(titulo).bold = True
                p.alignment = WD_ALIGN_PARAGRAPH.LEFT

                if i + 1 < len(doc.paragraphs):
                    proximo = doc.paragraphs[i + 1]

                    # Remove recuo do próximo parágrafo também
                    proximo.paragraph_format.left_indent = None
                    proximo.paragraph_format.first_line_indent = None

                    palavras = re.split(r"[;,.\n]\s*", proximo.text.strip())
                    palavras = [w.strip() for w in palavras if w.strip()]
                    proximo.text = "; ".join(palavras) + "."
                    proximo.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY

            # ➜ espaço 1,5 linha
            p.paragraph_format.space_after = Pt(18)
            return

def formatar_titulos_numerados(doc):
    import re
    from docx.shared import Pt
    from docx.enum.text import WD_ALIGN_PARAGRAPH

    padrao = r"""
        ^\s*
        (\d+(\.\d+)*)   # 1 , 1.1 , 1.1.1 etc
        \s*\.?\s*
        [A-Za-zÀ-ÿ]     # precisa começar com letra depois do número
    """
    eh_titulo = lambda t: re.match(padrao, t.strip(), re.VERBOSE) is not None

    depois_do_resumo = False

    for p in doc.paragraphs:
        texto_original = p.text.strip()

        if not texto_original:
            continue

        if re.match(r"^\s*resumo\b", texto_original, re.IGNORECASE):
            depois_do_resumo = True
            continue

        if not depois_do_resumo:
            continue

        if eh_titulo(texto_original):

            print(f"[DEBUG] Título numerado detectado: {texto_original}")

            match = re.match(r"^(\d+(?:\.\d+)*)(?:\.)?\s*(.*)$", texto_original)
            if not match:
                continue

            numeracao = match.group(1)
            titulo_texto = match.group(2)

            nivel = numeracao.count(".") + 1

            if nivel == 1:
                titulo_formatado = titulo_texto.upper()
            else:
                if len(titulo_texto) > 0:
                    titulo_formatado = titulo_texto[0].upper() + titulo_texto[1:].lower()
                else:
                    titulo_formatado = titulo_texto

            novo_texto = f"{numeracao} {titulo_formatado}"

            p.text = novo_texto

            p.alignment = WD_ALIGN_PARAGRAPH.LEFT
            for run in p.runs:
                run.font.bold = True
                run.font.size = Pt(12)
                # cor removida

def formatar_paragrafos_abnt(doc):

    mensagens = []
    depois_do_resumo = False

    padrao_titulo = r"^\s*\d+(\.\d+)*\s*[\).]?\s*[A-Za-zÀ-ÿ]"

    def eh_titulo(t):
        return re.match(padrao_titulo, t.strip()) is not None

    for p in doc.paragraphs:
        texto = p.text.strip()
        if not texto:
            continue

        if re.match(r"^\s*resumo\b", texto, re.IGNORECASE):
            depois_do_resumo = True
            continue

        if not depois_do_resumo:
            continue

        if eh_titulo(texto):
            continue

        p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
        p.paragraph_format.first_line_indent = Pt(1.25 * 28.35)
        p.paragraph_format.space_after = Pt(6)
        p.paragraph_format.line_spacing = 1.5
        p.paragraph_format.right_indent = Pt(0)


        for run in p.runs:
            run.font.size = Pt(12)
            run.font.bold = False
            # cor removida

        erros = []

        if p.alignment != WD_ALIGN_PARAGRAPH.JUSTIFY:
            erros.append("alinhamento incorreto (esperado: justificado)")

        recuo = p.paragraph_format.first_line_indent
        if not recuo or abs(recuo.cm - 1.25) > 0.05:
            erros.append("recuo incorreto (esperado: 1,25 cm)")

        if p.paragraph_format.space_after != Pt(6):
            erros.append("espaçamento depois incorreto (esperado: 6pt)")

        if p.paragraph_format.line_spacing != 1.5:
            erros.append("espaçamento de linha incorreto (esperado: 1,5)")

        for run in p.runs:
            if run.font.size and run.font.size.pt != 12:
                erros.append("tamanho da fonte errado (esperado: 12)")
                break

        if erros:
            mensagens.append(f"❌ Parágrafo incorreto: \"{texto[:50]}...\" → " + " | ".join(erros))
        else:
            mensagens.append(f"✅ Parágrafo OK: \"{texto[:50]}...\"")

    return mensagens

def formatar_referencias(doc):
    import re
    from docx.shared import Pt
    from docx.enum.text import WD_ALIGN_PARAGRAPH
    import unicodedata

    try:
        normalizar = lambda s: remover_acentos(s or "").strip().lower()
    except NameError:
        def normalizar(s):
            if not s:
                return ""
            s = ''.join(c for c in unicodedata.normalize('NFD', s)
                        if unicodedata.category(c) != 'Mn')
            s = s.lower().strip()
            s = re.sub(r'\s+', ' ', s)
            return s

    padrao_titulo = re.compile(r"^\s*\d*\.?\s*REFERÊNCIAS\s*$", re.IGNORECASE)

    inicio_idx = None
    for i, p in enumerate(doc.paragraphs):
        texto_norm = " ".join(p.text.upper().split())
        if padrao_titulo.match(texto_norm):
            inicio_idx = i + 1
            break

    if inicio_idx is None:
        return

    def eh_titulo_numerado(texto):
        return bool(re.match(r"^\s*\d+(\.\d+)*\s+[A-Za-zÀ-ÿ]", texto))

    ano_re = re.compile(r"\b(19|20)\d{2}\b")

    def eh_referencia(t):
        if not t or len(t) < 8:
            return False

        tn = normalizar(t)
        pontos = 0

        if "http" in tn or "doi" in tn:
            pontos += 2

        if ano_re.search(t):
            pontos += 1

        palavras_comuns = [
            "disponivel em", "acesso em", "editora", "vol", "v.", "n.",
            "nbr", "issn", "isbn", "revista", "congresso", "artigo",
            "tecnologia da informação", "internet", "pesquisa", "relatório"
        ]
        if any(k in tn for k in palavras_comuns):
            pontos += 1

        if re.search(r"[A-Za-zÀ-ÿ ]+:\s*[A-Za-zÀ-ÿ ]+", t):
            pontos += 1

        if re.match(r"^[A-Z]{2,}\b", t):
            pontos += 1

        if re.match(r"^[A-ZÀ-Ý]{2,}\s*,", t):
            pontos += 1

        return pontos >= 1.5

    for p in doc.paragraphs[inicio_idx:]:

        texto = p.text.strip()
        if not texto:
            continue

        if eh_titulo_numerado(texto):
            break

        if eh_referencia(texto):

            p.alignment = WD_ALIGN_PARAGRAPH.LEFT
            p.paragraph_format.line_spacing = 1.0
            p.paragraph_format.space_after = Pt(6)
            p.paragraph_format.right_indent = Pt(0)
            p.paragraph_format.first_line_indent = Pt(0)
            p.paragraph_format.left_indent = Pt(1.25 * 28.35)

            for run in p.runs:
                run.font.size = Pt(12)
                run.font.bold = False
                # cor removida



# -------- APLICAR FORMATAÇÃO --------
from docx.shared import Pt, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
import re

def aplicar_formatacao(doc, fonte_principal):
    """
    Aplica formatação visual (temporária/colorida) aos elementos da CAPA
    antes do resumo, apenas para facilitar depuração.

    Itens identificados:
        - instituição  (verde)
        - autor        (azul)
        - orientador   (magenta)
        - curso        (laranja)
        - título       (verde escuro)
        - subtítulo    (roxo)
        - cidade       (ciano)
        - ano          (cinza)
    """

    # ==============================================
    # 1) Garantir fonte padrão
    # ==============================================
    for p in doc.paragraphs:
        for run in p.runs:
            run.font.name = fonte_principal
            try:
                rFonts = run._element.rPr.rFonts
                rFonts.set(qn("w:ascii"), fonte_principal)
                rFonts.set(qn("w:hAnsi"), fonte_principal)
            except:
                pass

    antes_do_resumo = True
    titulo_identificado = False

    # ==============================================
    # 2) Flags para cada item da capa
    # ==============================================
    encontrado = {k: False for k in [
        "instituicao", "autor", "orientador", "curso",
        "titulo_capa", "subtitulo", "cidade", "ano"
    ]}

    # Paleta para depuração

    # ==============================================
    # 3) Percorre documento até o resumo
    # ==============================================
    for p in doc.paragraphs:

        texto = p.text.strip()
        if not texto:
            continue

        # Parou ao encontrar o resumo
        if re.match(r"^\s*resumo\b", texto, re.IGNORECASE):
            antes_do_resumo = False
            continue

        if not antes_do_resumo:
            continue

        # Classificação
        tipo, titulo_identificado = classificar_texto(texto, titulo_identificado)
        print(f"[DEBUG] Linha: '{texto}'")
        print(f"[DEBUG] Tipo detectado: {tipo}")

        # Se for algo não reconhecido, ignora
        if tipo not in encontrado:
            continue

        # Garante que apenas o primeiro de cada tipo será formatado
        if encontrado[tipo]:
            continue

        encontrado[tipo] = True

        # ==========================================
        # 4) Aplicação da cor e estilo
        # ==========================================
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER


        # Ajustes especiais
        if tipo == "titulo_capa":
            p.text = texto.upper()
            for run in p.runs:
                run.font.size = Pt(14)
                run.bold = True

        elif tipo == "cidade":
            for run in p.runs:
                run.text = " ".join(w.capitalize() for w in run.text.lower().split())

        elif tipo == "orientador":
            for run in p.runs:
                run.text = run.text.replace(":", "").strip().capitalize()

        print(f"[DEBUG] Formatação aplicada para: {tipo}")

    # Fim da função

def aplicar_margens_abnt(doc):
    s = doc.sections[0]
    s.top_margin = Cm(3)
    s.bottom_margin = Cm(2)
    s.left_margin = Cm(3)
    s.right_margin = Cm(2)

  
# -------- VERIFICAR FORMATAÇÃO --------
def verificar_formatacao(doc):
    from docx.enum.text import WD_ALIGN_PARAGRAPH
    import re

    erros = set()

    # --------------------------------------------
    # 1. Padrão de detecção de TÍTULOS NUMERADOS
    # --------------------------------------------
    padrao_titulo = r"""
        ^\s*
        (\d+(\.\d+)*)       # 1 , 1.1 , 1.1.1 etc
        \s*\.?\s*
        [A-Za-zÀ-ÿ]         # letra após número
    """
    eh_titulo_numerado = lambda t: re.match(padrao_titulo, t.strip(), re.VERBOSE) is not None

    # --------------------------------------------
    # 2. Flags
    # --------------------------------------------
    dentro_corpo = False  # ativa somente após encontrar "RESUMO"

    # --------------------------------------------
    # 3. Percorre todo o documento
    # --------------------------------------------
    for p in doc.paragraphs:
        texto = p.text.strip()
        if not texto:
            continue

        # Detecta INÍCIO do corpo (após RESUMO)
        if re.match(r"^\s*resumo\b", texto, re.IGNORECASE):
            dentro_corpo = True
            continue

        # Somente após o RESUMO
        if not dentro_corpo:
            continue

        # --------------------------------------------
        # 4. Sessão detectada (ex: 1 Introdução)
        # --------------------------------------------
        if eh_titulo_numerado(texto):

            # Separa numeração e conteúdo
            match = re.match(r"^(\d+(?:\.\d+)*)(?:\.)?\s*(.*)$", texto)
            if not match:
                continue

            numeracao = match.group(1)
            titulo_texto = match.group(2)

            # Nível do título (1, 2, 3 etc)
            nivel = numeracao.count(".") + 1

            # — Tamanho da fonte
            tamanhos = {run.font.size.pt for run in p.runs if run.font.size}
            if any(t != 12 for t in tamanhos):
                erros.add(f"❌ A sessão '{texto}' deve estar no tamanho 12.")

            # — Alinhamento correto
            if p.alignment not in (WD_ALIGN_PARAGRAPH.LEFT, None):
                erros.add(f"⚠️ A sessão '{texto}' deve estar alinhada à esquerda.")

            # — Regras de maiúscula
            if nivel == 1 and titulo_texto != titulo_texto.upper():
                erros.add(f"⚠️ Sessão de nível 1 '{texto}' deveria estar em MAIÚSCULO.")

            if nivel > 1 and titulo_texto == titulo_texto.upper():
                erros.add(f"⚠️ Sessões de nível 2+ não devem ser totalmente maiúsculas: '{texto}'.")

            continue  # evita cair na análise de parágrafo

        # --------------------------------------------
        # 5. Parágrafos normais (somente após RESUMO)
        # --------------------------------------------

        # Verificar fonte
        tamanhos = {run.font.size.pt for run in p.runs if run.font.size}
        if any(t != 12 for t in tamanhos):
            erros.add("❌ Um parágrafo está no tamanho incorreto (deve ser 12).")

        # Verificar alinhamento — mas ignorar None
        if p.alignment not in (WD_ALIGN_PARAGRAPH.JUSTIFY, None):
            erros.add("⚠️ Um ou mais parágrafos não estão justificados.")

    # --------------------------------------------
    # 6. Retorno final
    # --------------------------------------------
    if not erros:
        return ["✅ Nenhum erro encontrado na formatação pós-resumo."]

    return list(erros)


def verificar_margens(doc):
    padrao = [3, 2, 3, 2]
    s = doc.sections[0]
    margens = [round(s.top_margin.cm), round(s.bottom_margin.cm),
               round(s.left_margin.cm), round(s.right_margin.cm)]
    return margens == padrao


@app.route("/formatar", methods=["POST"])
def formatar():
    if "arquivo" not in request.files:
        return {"erro": "Envie um arquivo .docx"}, 400

    arquivo = request.files["arquivo"]
    caminho = tempfile.NamedTemporaryFile(delete=False, suffix=".docx").name
    arquivo.save(caminho)

    doc = Document(caminho)
    fonte = detectar_fonte_principal(doc)

    # -------- AJUSTA ESTILO NORMAL --------
    try:
        estilo_normal = doc.styles["Normal"].font
        estilo_normal.name = fonte
        for r in doc.styles["Normal"].element.xpath(".//w:rFonts"):
            r.set(qn("w:ascii"), fonte)
            r.set(qn("w:hAnsi"), fonte)
    except:
        pass

    formatar_capa(doc)
    formatar_paragrafos_abnt(doc)
    formatar_resumo(doc)
    formatar_palavras_chave(doc)
    formatar_abstract(doc) 
    formatar_keywords(doc) 
    formatar_titulos_numerados(doc) 
    formatar_referencias(doc)
    aplicar_margens_abnt(doc)
    aplicar_formatacao(doc, fonte)

    saida = tempfile.NamedTemporaryFile(delete=False, suffix=".docx").name
    doc.save(saida)
    return send_file(saida, as_attachment=True, download_name="arquivo_formatado_ABNT.docx")


@app.route("/verificar", methods=["POST"])
def verificar():
    if "arquivo" not in request.files:
        return {"erro": "Envie um arquivo .docx"}, 400

    arquivo = request.files["arquivo"]
    caminho = tempfile.NamedTemporaryFile(delete=False, suffix=".docx").name
    arquivo.save(caminho)
    doc = Document(caminho)

    return {
        "margens_corretas": verificar_margens(doc),
        "formatacao": verificar_formatacao(doc)
    }


if __name__ == "__main__":
    app.run(debug=True)
